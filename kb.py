"""Curated knowledge base — user-uploadable per-app docs (T3-2).

The user's mental model: put documentation for an app in ``KB_DIR``, named to match the
``.exe`` basename, and Nimbus treats it as authoritative on the next push-to-talk.

Two layouts, and **both work**:

* **Flat file** — ``KB_DIR/myapp.exe.md``. The original design and still the simplest thing
  that works. Unchanged.
* **Folder** — ``KB_DIR/myapp.exe/`` containing any number of ``.md``, ``.txt``, ``.pdf`` or
  ``.docx`` files, read in stable alphabetical order. Added by T3-2.

If both exist the flat file is read **first**, so an existing user who later adds a folder
keeps their original file as the leading context rather than having it reordered.

Missing entirely → ``recall()`` returns ``('', '')``, the pipeline skips KB injection, and
Nimbus answers from vision plus memory. That is the "Nimbus already knows that software" path
and it is the common case.

## Why this matters most for in-house software

Public tools are already in the model's training data. A company's internal CAD tool,
ticketing system or trading terminal is not, and no amount of screen-reading recovers a
convention that only exists in a wiki. This is the only mechanism that closes that gap, which
is why T3-2 raised the ceiling from one Markdown file to a folder of real documents.

## Over-budget behaviour changed, deliberately

``KB_RECALL_MAX_CHARS`` is 60,000. Previously an over-budget file was **tail-truncated** —
Nimbus silently read the last 60k and discarded everything before it, so a question about
anything in the discarded part was answered from nothing.

When a query is supplied, sections are now **ranked by keyword overlap** and the best kept
until the budget fills. Callers that pass no query keep the exact previous behaviour.

Ranking is keyword-based, not embeddings. ``memory.py`` argues against premature vector-DB
complexity and that judgement stands (§8): plain files the user can read, edit and delete are
the transparency contract, and a keyword score is inspectable in a way a cosine distance is
not.

## Deviation from the original plan, and why

The audit specified PDFs via the Gemini **Files API** (``T1-6c``). That was not built, and
should not be: ``kb_content`` is a **string** injected into the system prompt and flows
provider-agnostically through the whole pipeline. Routing one file format through a
Gemini-only file-reference path would fracture that contract and break PDFs on the
fully-local Ollama path, which §1.6 lists as a regression gate. Local text extraction keeps
one code path for five providers and stays inspectable. ``T1-6c`` is therefore unnecessary
rather than outstanding.

Extractors are **lazily imported and individually optional**: a missing ``pypdf`` skips PDFs
with a log rather than breaking the knowledge base, mirroring how ``faster_whisper`` and
``kokoro_onnx`` degrade.
"""
from __future__ import annotations

import re
from pathlib import Path

from config import KB_DIR, KB_RECALL_MAX_CHARS
from memory import _sanitize_app_name as _memory_sanitize


TEXT_SUFFIXES: frozenset[str] = frozenset({".md", ".markdown", ".txt", ".text"})
"""Suffixes read directly as UTF-8."""

PDF_SUFFIXES: frozenset[str] = frozenset({".pdf"})
DOCX_SUFFIXES: frozenset[str] = frozenset({".docx"})

SUPPORTED_SUFFIXES: frozenset[str] = TEXT_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES
"""Everything ``recall`` will attempt. Anything else in the folder is ignored silently --
users keep images and spreadsheets alongside their notes and that must not be an error."""

MAX_FILES_PER_APP = 40
"""Cap on files read from one app folder.

A guard against a user pointing Nimbus at a directory with thousands of files: reading them
all would stall the interaction on the hot path. The cap applies before extraction, so the
cost is bounded regardless of file sizes."""

_MIN_TERM_LENGTH = 3
"""Shortest query word used for ranking. Drops "is", "a", "to" -- they match everything and
therefore rank nothing."""

_SECTION_SPLIT_RE = re.compile(r"\n(?=#{1,6}\s)")
"""Split Markdown on headings. Sections are the natural ranking unit: a heading plus its body
is a self-contained topic, which is what a question is usually about."""


GUIDE_FILENAME = "README.md"
"""Name of the self-seeded guide written into ``KB_DIR``.

Safe from being mistaken for knowledge-base content: a flat lookup only ever reads
``<app>.md``, so this would need a foreground application literally named ``README`` with no
file extension. ``get_foreground_app()`` returns an ``.exe`` basename, so that cannot occur.
"""

_GUIDE_TEXT = """\
# Nimbus knowledge base

Drop documentation here and Nimbus treats it as authoritative for that application.

This is most useful for software the model has **never seen** - an internal tool, a
company-specific workflow, a plugin with no public documentation. Nimbus already knows Excel
and Blender; it knows nothing about your in-house scheduler.

## Naming

Files are matched to the application's `.exe` name, lowercased.

To find the exact name for an app, use it once with Nimbus and then look in your memory
folder - the filename there is the name to use:

```
%USERPROFILE%\\.nimbus\\memory\\
```

For example, if you see `orionflow.exe.md` there, the app's name is `orionflow.exe`.

## Two layouts

**One file** - simplest, good for a page or two of notes:

```
Nimbus Wiki\\
    orionflow.exe.md
```

**A folder** - for anything larger, or for real documents:

```
Nimbus Wiki\\
    orionflow.exe\\
        01-overview.md
        02-shortcuts.docx
        03-troubleshooting.pdf
        guides\\
            exporting.md
```

Both work, and both can be used at once. Files are read in alphabetical order, so a numeric
prefix (`01-`, `02-`) controls what Nimbus reads first. Subfolders are included.

## Supported formats

| Format | Notes |
|---|---|
| `.md`, `.markdown` | Headings help - see below |
| `.txt`, `.text` | Read as-is |
| `.pdf` | Text is extracted; scanned images are not read |
| `.docx` | Paragraphs **and tables** - good for shortcut lists |

Anything else (images, spreadsheets, videos) is ignored, so you can keep them alongside your
notes without causing an error.

## Use headings

When your documentation is larger than Nimbus can send in one request (60,000 characters,
roughly 15 pages), it selects the sections most relevant to your question rather than
truncating blindly.

That selection works on Markdown headings, so this:

```markdown
# Exporting
...

# Troubleshooting
...
```

lets Nimbus send just the export section when you ask about exporting. Without headings the
whole file is treated as one block, and a large file may get cut.

## Limits

- 60,000 characters sent per question, selected by relevance
- 40 files per application

## Notes

- Changes are picked up on your next question. No restart needed.
- Nothing here is uploaded anywhere except as part of a question you ask, to whichever model
  provider you have configured. On a local provider (Ollama) it never leaves your machine.
- These are ordinary files. Edit or delete them freely.
- Nimbus rewrites this guide if it is missing. Your own files are never touched.
"""
"""The guide, embedded as a string rather than shipped as a data file.

Deliberate: a ``datas`` entry in ``nimbus.spec`` would need ``sys._MEIPASS`` path resolution
that differs between the frozen build and a source checkout, which is a recurring source of
"works in dev, missing in the installer" bugs. 3 KB of text in the module has none of that
failure mode and cannot go missing.
"""


def ensure_guide(kb_dir: Path | None = None) -> Path | None:
    """Write the guide into ``KB_DIR`` if absent. Returns the path, or ``None`` on failure.

    **Why this exists.** ``config._resolve_kb_dir`` creates the folder at startup, so a new
    user gets an empty ``Documents\\Nimbus Wiki`` with no indication that a file must be named
    ``orionflow.exe.md`` to match the executable. That convention is not guessable, so without
    a guide the entire feature is invisible and effectively unused.

    Only writes when the file does not exist, so a user's own edits are never overwritten.
    Never raises: a read-only or redirected folder must not prevent Nimbus starting, and the
    guide is help text, not functionality.
    """
    base = Path(kb_dir) if kb_dir is not None else Path(KB_DIR)
    path = base / GUIDE_FILENAME
    try:
        if path.exists():
            return path
        base.mkdir(parents=True, exist_ok=True)
        path.write_text(_GUIDE_TEXT, encoding="utf-8")
        return path
    except OSError:
        return None


def _sanitize_app_name(app_name: str) -> str:
    """Normalise an app name to its KB filename stem.

    **Delegates to ``memory._sanitize_app_name`` rather than reimplementing it (T3-2).**
    This module previously had its own copy whose docstring claimed to mirror memory's
    "exactly". It did not, and measurement found 7 of 15 test inputs disagreeing: memory
    strips surrounding whitespace and replaces all nine Windows-reserved characters, while
    this copy stripped nothing and replaced only three.

    That broke the documented mental model. Users are told to read the canonical name out of
    ``~/.nimbus/memory/`` and name their KB file to match -- but for an app whose name needed
    stripping, memory showed ``spaced.exe`` while this function looked for ``  spaced.exe  ``
    and found nothing. One function means the two folders can never disagree again.

    Returns ``""`` for input memory rejects, since ``recall`` treats empty as "no KB".
    """
    try:
        return _memory_sanitize(app_name)
    except ValueError:
        return ""


# --- Extraction --------------------------------------------------------------

def extract_pdf_text(path: Path) -> str:
    """Extract text from a PDF. Returns ``""`` if ``pypdf`` is unavailable."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_docx_text(path: Path) -> str:
    """Extract text from a .docx, including table cells.

    Tables are included because software documentation puts keyboard shortcuts, field
    definitions and option matrices in them -- exactly the content worth having.
    """
    try:
        import docx
    except ImportError:
        return ""
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    """Extract text from any supported file. Returns ``""`` on anything unreadable.

    Never raises. KB files are user-controlled -- a corrupt PDF, an exotic encoding or a
    file locked by another program must not break the interaction. ``app.py`` already wraps
    ``recall`` in try/except; this keeps one bad file from costing the user the *rest* of
    their knowledge base too.
    """
    suffix = path.suffix.lower()
    try:
        if suffix in TEXT_SUFFIXES:
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix in PDF_SUFFIXES:
            return extract_pdf_text(path)
        if suffix in DOCX_SUFFIXES:
            return extract_docx_text(path)
    except Exception:
        return ""
    return ""


def iter_kb_files(folder: Path) -> list[Path]:
    """Supported files under ``folder``, recursively, in stable order.

    Sorted by relative POSIX path so the concatenation order is identical on every run.
    Unstable ordering would silently change what survives truncation between two otherwise
    identical questions, which is the kind of non-determinism that makes a bug
    unreproducible.
    """
    if not folder.is_dir():
        return []
    files = [
        p for p in folder.rglob("*")
        if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES
    ]
    files.sort(key=lambda p: p.relative_to(folder).as_posix().lower())
    return files[:MAX_FILES_PER_APP]


# --- Relevance ranking -------------------------------------------------------

def query_terms(query: str) -> set[str]:
    """Content words from a query, lowercased.

    Words shorter than ``_MIN_TERM_LENGTH`` are dropped: they match nearly everything and so
    contribute no ranking signal.
    """
    return {
        word for word in re.findall(r"[a-z0-9]+", (query or "").lower())
        if len(word) >= _MIN_TERM_LENGTH
    }


def split_sections(text: str) -> list[str]:
    """Split text into ranking units, on Markdown headings where present.

    Falls back to the whole text as one section when there are no headings -- a plain
    extracted PDF has none, and splitting it arbitrarily would cut mid-sentence.
    """
    if not text.strip():
        return []
    sections = [s for s in _SECTION_SPLIT_RE.split(text) if s.strip()]
    return sections or [text]


def score_section(section: str, terms: set[str]) -> int:
    """How many distinct query terms appear in ``section``.

    Distinct terms, not total occurrences: a section repeating one word fifty times is not
    more relevant than one covering five of the query's words. Counting occurrences would let
    a glossary entry outrank the page that actually answers the question.
    """
    if not terms:
        return 0
    lowered = section.lower()
    return sum(1 for term in terms if term in lowered)


def rank_and_truncate(text: str, query: str, max_chars: int) -> str:
    """Fit ``text`` into ``max_chars``, keeping the most relevant sections (T3-2).

    Under budget, returns the text unchanged -- ranking must not reorder content that all
    fits, because a document's own order carries meaning.

    Over budget with no usable query, falls back to **tail-truncation**, preserving the
    previous behaviour exactly.

    Over budget with a query, keeps the highest-scoring sections but re-emits them in their
    **original document order**, so headings still read in sequence rather than by score.
    """
    if len(text) <= max_chars:
        return text
    terms = query_terms(query)
    sections = split_sections(text)
    if not terms or len(sections) <= 1:
        return text[-max_chars:]

    scored = [
        (score_section(section, terms), index, section)
        for index, section in enumerate(sections)
    ]
    # Highest score first; earlier sections win ties, since a document's opening usually
    # carries the overview.
    scored.sort(key=lambda item: (-item[0], item[1]))

    kept: list[tuple[int, str]] = []
    used = 0
    for score, index, section in scored:
        if score == 0 and kept:
            # Once relevant sections are in, irrelevant ones are not worth budget.
            break
        cost = len(section) + 1
        if used + cost > max_chars:
            continue
        kept.append((index, section))
        used += cost
    if not kept:
        return text[-max_chars:]
    kept.sort(key=lambda pair: pair[0])
    return "\n".join(section for _, section in kept)[:max_chars]


# --- Public API --------------------------------------------------------------

def recall(
    app_name: str,
    kb_dir: Path | None = None,
    max_chars: int = KB_RECALL_MAX_CHARS,
    query: str = "",
) -> tuple[str, str]:
    """Look up the curated KB for ``app_name``.

    Args:
        app_name: foreground .exe basename, e.g. ``"MYAPP.EXE"``.
        kb_dir: override the KB folder (test hook). Defaults to ``config.KB_DIR``.
        max_chars: budget. Defaults to ``config.KB_RECALL_MAX_CHARS``.
        query: the user's transcript. When supplied AND the content exceeds the budget,
            sections are ranked by keyword overlap instead of blindly tail-truncated.
            Omitting it reproduces the original behaviour exactly.

    Returns:
        ``(content, sanitized_name)``, or ``("", "")`` when nothing matches, ``app_name`` is
        blank, or ``max_chars <= 0``.

        That last guard is not defensive noise: Python's ``text[-0:]`` returns the WHOLE
        string, so a caller passing 0 from a misconfigured override would get the full file
        rather than none of it.
    """
    if max_chars <= 0:
        return ("", "")
    if not app_name or not app_name.strip():
        return ("", "")

    sanitized = _sanitize_app_name(app_name)
    if not sanitized:
        return ("", "")

    base_dir = Path(kb_dir) if kb_dir is not None else Path(KB_DIR)
    parts: list[str] = []

    # Flat file first, so an existing user who later adds a folder keeps their original
    # file as the leading context.
    flat = base_dir / f"{sanitized}.md"
    if flat.is_file():
        text = extract_text(flat)
        if text.strip():
            parts.append(text)

    folder = base_dir / sanitized
    for path in iter_kb_files(folder):
        text = extract_text(path)
        if not text.strip():
            continue
        # Name each file inline. The model benefits from knowing that a shortcut list and a
        # troubleshooting guide are separate documents, and the user can tell from Nimbus's
        # answer which of their files it came from.
        relative = path.relative_to(folder).as_posix()
        parts.append(f"## {relative}\n\n{text}")

    if not parts:
        return ("", "")

    combined = "\n\n".join(parts)
    return (rank_and_truncate(combined, query, max_chars), sanitized)


if __name__ == "__main__":
    print("=" * 70)
    print("Nimbus -- kb.py manual verification")
    print(f"  KB_DIR: {KB_DIR}")
    print(f"  KB_RECALL_MAX_CHARS: {KB_RECALL_MAX_CHARS}")
    print(f"  supported: {sorted(SUPPORTED_SUFFIXES)}")
    print("=" * 70)

    for app in ("MYAPP.EXE", "myapp.exe", "Fusion360.exe", "DOES_NOT_EXIST.EXE"):
        content, name = recall(app)
        if content:
            preview = content[:120].replace("\n", " ")
            print(f"\n{app:30s} -> matched {name!r}, {len(content)} chars")
            print(f"  preview: {preview!r}...")
        else:
            print(f"\n{app:30s} -> nothing for {_sanitize_app_name(app)!r}")

    print(f"\nFlat file:  {KB_DIR / 'myapp.exe.md'}")
    print(f"Or folder:  {KB_DIR / 'myapp.exe'}{chr(92)}  (.md .txt .pdf .docx)")
